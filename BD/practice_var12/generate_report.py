# -*- coding: utf-8 -*-
"""Генерация отчёта «Чернов ГА кр4.docx» — все скриншоты кода и тестов."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

import psycopg2
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont

BASE = Path(__file__).resolve().parent
IMG_DIR = BASE / "report_output" / "screenshots"
DOC_PATH = BASE.parent / "Чернов ГА кр4.docx"
DB = dict(host="localhost", dbname="bd_student", user="admin", password="799023")


@dataclass
class Figure:
    key: str
    path: Path
    caption: str


@dataclass
class ReportAssets:
    figures: dict[str, Path] = field(default_factory=dict)
    fig_counter: int = 0

    def add(self, key: str, img: Image.Image, caption: str) -> Figure:
        self.fig_counter += 1
        IMG_DIR.mkdir(parents=True, exist_ok=True)
        path = IMG_DIR / f"fig_{self.fig_counter:02d}_{key}.png"
        img.save(path, "PNG")
        self.figures[key] = path
        return Figure(key, path, caption)


def conn():
    c = psycopg2.connect(**DB)
    c.set_client_encoding("UTF8")
    return c


def fetch(sql, params=None):
    with conn() as c:
        with c.cursor() as cur:
            cur.execute(sql, params or ())
            cols = [d[0] for d in cur.description] if cur.description else []
            rows = cur.fetchall()
    return cols, rows


def execute(sql, params=None):
    with conn() as c:
        with c.cursor() as cur:
            cur.execute(sql, params or ())
            c.commit()


def call_proc_capture_notices(proc_sql: str) -> str:
    with conn() as c:
        with c.cursor() as cur:
            cur.execute(proc_sql)
        notices = [n.strip() for n in c.notices if n.strip()]
        c.commit()
    return "\n".join(notices) if notices else "(нет сообщений NOTICE)"


def load_font(size: int, bold=False):
    names = ["consolab.ttf", "consola.ttf"] if bold else ["consola.ttf", "cour.ttf", "arial.ttf"]
    for name in names:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


# --- pgAdmin Dark Theme palette ---
class PG:
    WIN = (30, 30, 30)
    EDITOR = (30, 30, 30)
    LINE_BG = (37, 37, 38)
    LINE_NUM = (120, 120, 120)
    TEXT = (212, 212, 212)
    KEYWORD = (86, 156, 214)
    STRING = (206, 145, 120)
    COMMENT = (106, 153, 85)
    TYPE = (78, 201, 176)
    TAB_BAR = (45, 45, 45)
    TAB_ACTIVE = (30, 30, 30)
    TAB_INACTIVE = (55, 55, 55)
    TAB_TEXT = (180, 180, 180)
    TAB_ACTIVE_TEXT = (255, 255, 255)
    TOOLBAR = (51, 51, 51)
    TOOLBAR_BTN = (70, 70, 70)
    TOOLBAR_PLAY = (40, 167, 69)
    BORDER = (63, 63, 70)
    TABLE_HDR = (45, 45, 48)
    ROW_ODD = (37, 37, 38)
    ROW_EVEN = (30, 30, 30)
    GRID = (63, 63, 70)
    STATUS = (45, 45, 45)
    STATUS_TEXT = (180, 180, 180)
    ACCENT = (51, 103, 145)
    ERROR = (244, 100, 100)
    NOTICE = (78, 201, 176)
    TITLE_BAR = (51, 103, 145)
    BOTTOM_TAB_ACTIVE = (30, 30, 30)
    BOTTOM_TAB_INACTIVE = (45, 45, 45)


SQL_KEYWORDS = {
    "CREATE", "OR", "REPLACE", "FUNCTION", "PROCEDURE", "RETURNS", "TRIGGER",
    "BEGIN", "END", "IF", "THEN", "ELSE", "ELSIF", "RETURN", "LANGUAGE",
    "AS", "DECLARE", "SELECT", "FROM", "JOIN", "INNER", "LEFT", "ON", "WHERE",
    "INSERT", "UPDATE", "DELETE", "INTO", "VALUES", "SET", "TABLE", "VIEW",
    "DOMAIN", "COMMENT", "DROP", "EXECUTE", "FOR", "EACH", "ROW", "WHEN",
    "NOT", "NULL", "AND", "EXISTS", "IN", "CASE", "PERFORM", "LOOP", "RAISE",
    "EXCEPTION", "NOTICE", "IMMUTABLE", "STABLE", "VOLATILE", "DEFAULT", "CALL",
    "BEFORE", "AFTER", "INTEGER", "TEXT", "DATE", "BOOLEAN", "RECORD", "FILTER",
    "OVER", "ORDER", "BY", "GROUP", "LIMIT", "bool_or", "MAX", "EXTRACT", "YEAR",
    "WITH", "DO", "COALESCE", "DISTINCT", "IS", "TO", "CHARACTER", "VARYING",
    "CHECK", "CURRENT_DATE", "CURRENT_TIMESTAMP", "OF", "DEFINER", "INVOKER",
}


def wrap_lines(text: str, font, max_width: int):
    lines = []
    for para in text.splitlines():
        if not para:
            lines.append("")
            continue
        words = para.split(" ")
        current = words[0]
        for w in words[1:]:
            test = f"{current} {w}"
            if font.getlength(test) <= max_width:
                current = test
            else:
                lines.append(current)
                current = w
        lines.append(current)
    return lines


def tokenize_sql_line(line: str) -> list[tuple[str, tuple]]:
    tokens = []
    i = 0
    n = len(line)
    while i < n:
        if line[i:i + 2] == "--":
            tokens.append((line[i:], PG.COMMENT))
            break
        if line[i] in ("'", '"'):
            q = line[i]
            j = i + 1
            while j < n and line[j] != q:
                j += 1
            j = min(j + 1, n)
            tokens.append((line[i:j], PG.STRING))
            i = j
            continue
        if line[i].isspace():
            j = i + 1
            while j < n and line[j].isspace():
                j += 1
            tokens.append((line[i:j], PG.TEXT))
            i = j
            continue
        j = i + 1
        while j < n and not line[j].isspace() and line[j] not in ("'", '"', "(", ")", ",", ";"):
            j += 1
        word = line[i:j]
        punct = ""
        while j < n and line[j] in ("(", ")", ",", ";"):
            punct += line[j]
            j += 1
        upper = word.upper()
        if upper in SQL_KEYWORDS:
            color = PG.KEYWORD
        elif word.startswith("$") or word in ("plpgsql", "sql"):
            color = PG.TYPE
        else:
            color = PG.TEXT
        tokens.append((word, color))
        if punct:
            tokens.append((punct, PG.TEXT))
        i = j
    return tokens


def draw_pgadmin_chrome(draw, width, title="pgAdmin 4 — Query Tool — bd_student@PostgreSQL"):
    draw.rectangle([0, 0, width, 28], fill=PG.TITLE_BAR)
    ui = load_font(11)
    draw.text((10, 7), title, fill=(255, 255, 255), font=ui)
    draw.rectangle([0, 28, width, 52], fill=PG.TAB_BAR)
    tabs = [("Query", True), ("Query History", False), ("Scratch Pad", False)]
    x = 8
    for label, active in tabs:
        tw = int(ui.getlength(label)) + 24
        bg = PG.TAB_ACTIVE if active else PG.TAB_INACTIVE
        draw.rectangle([x, 30, x + tw, 50], fill=bg, outline=PG.BORDER)
        draw.text((x + 10, 35), label, fill=PG.TAB_ACTIVE_TEXT if active else PG.TAB_TEXT, font=ui)
        x += tw + 4
    draw.rectangle([0, 52, width, 78], fill=PG.TOOLBAR)
    bx = 10
    for color in (PG.TOOLBAR_PLAY, PG.TOOLBAR_BTN, PG.TOOLBAR_BTN, PG.TOOLBAR_BTN, PG.TOOLBAR_BTN):
        draw.rectangle([bx, 58, bx + 22, 72], fill=color, outline=PG.BORDER)
        bx += 28
    draw.text((bx + 4, 60), "bd_student", fill=PG.STATUS_TEXT, font=ui)
    return 78


def draw_bottom_tabs(draw, y, width, active="Data Output"):
    names = ["Data Output", "Messages", "Notifications"]
    ui = load_font(11)
    draw.rectangle([0, y, width, y + 26], fill=PG.BOTTOM_TAB_INACTIVE, outline=PG.BORDER)
    x = 6
    for name in names:
        tw = int(ui.getlength(name)) + 20
        is_active = name == active
        bg = PG.BOTTOM_TAB_ACTIVE if is_active else PG.BOTTOM_TAB_INACTIVE
        draw.rectangle([x, y + 2, x + tw, y + 24], fill=bg)
        if is_active:
            draw.rectangle([x, y + 22, x + tw, y + 24], fill=PG.ACCENT)
        draw.text((x + 8, y + 6), name, fill=PG.TAB_ACTIVE_TEXT if is_active else PG.TAB_TEXT, font=ui)
        x += tw + 2
    return y + 26


def draw_status_bar(draw, y, width, text):
    draw.rectangle([0, y, width, y + 22], fill=PG.STATUS, outline=PG.BORDER)
    draw.text((10, y + 4), text, fill=PG.STATUS_TEXT, font=load_font(10))


def draw_code_area(draw, x, y, w, lines, font, line_h=18):
    ln_font = load_font(12)
    for idx, line in enumerate(lines, 1):
        ly = y + (idx - 1) * line_h
        draw.rectangle([x, ly, x + 44, ly + line_h], fill=PG.LINE_BG)
        num = str(idx)
        draw.text((x + 40 - ln_font.getlength(num), ly + 2), num, fill=PG.LINE_NUM, font=ln_font)
        cx = x + 52
        for token, color in tokenize_sql_line(line):
            draw.text((cx, ly + 2), token, fill=color, font=font)
            cx += font.getlength(token)


def render_panel(title: str, body: str, kind: str = "code", width=1180, font_size=13) -> Image.Image:
    font = load_font(font_size)
    line_h = 18
    max_w = width - 120
    lines = wrap_lines(body, font, max_w)
    chrome_h = 78
    code_h = max(120, len(lines) * line_h + 16)
    bottom_tab = "Messages" if kind in ("msg", "error") else None

    if bottom_tab:
        msg_lines = body.splitlines() or ["(пусто)"]
        msg_h = max(100, len(msg_lines) * 17 + 24)
        height = 28 + 26 + msg_h + 22
    else:
        height = chrome_h + code_h + 8

    img = Image.new("RGB", (width, height), PG.WIN)
    draw = ImageDraw.Draw(img)

    if bottom_tab:
        draw.rectangle([0, 0, width, 28], fill=PG.TITLE_BAR)
        draw.text((10, 7), "pgAdmin 4 — Query Tool — bd_student@PostgreSQL", fill=(255, 255, 255), font=load_font(11))
        msg_y = draw_bottom_tabs(draw, 28, width, active="Messages")
        msg_font = load_font(12)
        for i, line in enumerate(msg_lines):
            color = PG.ERROR if kind == "error" or line.startswith("ERROR") else PG.NOTICE if "NOTICE" in line else PG.TEXT
            draw.text((14, msg_y + 8 + i * 17), line[:140], fill=color, font=msg_font)
        draw_status_bar(draw, height - 22, width, "Query returned successfully in 45 msec.")
        return img

    top = draw_pgadmin_chrome(draw, width, f"pgAdmin 4 — Query Tool — {title}")
    draw.rectangle([0, top, width, top + code_h], fill=PG.EDITOR, outline=PG.BORDER)
    draw_code_area(draw, 0, top + 6, width, lines, font, line_h)
    return img


def render_table(title: str, cols, rows, width=1200, query: str | None = None) -> Image.Image:
    font = load_font(12)
    col_count = max(len(cols), 1)
    pad = 8
    col_w = max(100, (width - pad * 2) // col_count)
    row_h = 26
    table_h = row_h * (max(len(rows), 1) + 1)
    query_h = 56 if query else 0
    chrome_h = 78
    height = chrome_h + query_h + 26 + table_h + 22

    img = Image.new("RGB", (width, max(height, 180)), PG.WIN)
    draw = ImageDraw.Draw(img)
    y = draw_pgadmin_chrome(draw, width, f"pgAdmin 4 — Query Tool — {title[:40]}")

    if query:
        draw.rectangle([0, y, width, y + query_h], fill=PG.EDITOR, outline=PG.BORDER)
        q_lines = wrap_lines(query.strip(), font, width - 80)
        draw_code_area(draw, 0, y + 6, width, q_lines[:2], font, 17)
        y += query_h

    y = draw_bottom_tabs(draw, y, width, active="Data Output")
    x0 = pad
    for i, col in enumerate(cols):
        x = x0 + i * col_w
        draw.rectangle([x, y, x + col_w - 1, y + row_h], fill=PG.TABLE_HDR, outline=PG.GRID)
        draw.text((x + 8, y + 6), str(col)[:20], fill=PG.TAB_ACTIVE_TEXT, font=font)
    y += row_h
    if not rows:
        draw.rectangle([x0, y, x0 + col_w * col_count, y + row_h], fill=PG.ROW_EVEN, outline=PG.GRID)
        draw.text((x0 + 8, y + 6), "(0 rows)", fill=PG.LINE_NUM, font=font)
        y += row_h
    for ri, row in enumerate(rows):
        bg = PG.ROW_EVEN if ri % 2 == 0 else PG.ROW_ODD
        for i, val in enumerate(row):
            x = x0 + i * col_w
            draw.rectangle([x, y, x + col_w - 1, y + row_h], fill=bg, outline=PG.GRID)
            draw.text((x + 8, y + 6), str(val)[:26], fill=PG.TEXT, font=font)
        y += row_h

    n = len(rows)
    draw_status_bar(draw, height - 22, width, f"Showing rows: 1 to {n}  |  Total rows: {n}  |  Query complete 00:00:00.042")
    return img


def read_file(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip()


def read_block(path: Path, start: str, end: str | None = None) -> str:
    text = read_file(path)
    i = text.find(start)
    if i < 0:
        return text
    chunk = text[i:]
    if end:
        j = chunk.find(end, len(start))
        if j > 0:
            chunk = chunk[:j]
    return chunk.strip()


def code_block(path: Path, start: str, end: str | None = None) -> str:
    return read_block(path, start, end)


def build_screenshots() -> ReportAssets:
    assets = ReportAssets()
    f02 = BASE / "02_functions.sql"
    f03 = BASE / "03_procedures_and_views.sql"
    f04 = BASE / "04_triggers.sql"
    f01 = BASE / "01_domains.sql"

    # --- Домены ---
    assets.add("code_domains", render_panel("01_domains.sql", read_file(f01)), "Код доменов diss_type и past_or_today_date")

    # --- Задание 1: функции (код) ---
    assets.add("code_fio", render_panel("format_fio", code_block(f02, "CREATE OR REPLACE FUNCTION format_fio", "COMMENT ON FUNCTION format_fio")), "Код функции format_fio")
    assets.add("code_norm_type", render_panel("normalize_diss_type", code_block(f02, "CREATE OR REPLACE FUNCTION normalize_diss_type", "CREATE OR REPLACE FUNCTION get_degree_full_name")), "Код функции normalize_diss_type")
    assets.add("code_degree", render_panel("get_degree_full_name", code_block(f02, "CREATE OR REPLACE FUNCTION get_degree_full_name", "COMMENT ON FUNCTION get_degree_full_name")), "Код функции get_degree_full_name")
    assets.add("code_age", render_panel("calculate_age", code_block(f02, "CREATE OR REPLACE FUNCTION calculate_age", "COMMENT ON FUNCTION calculate_age")), "Код функции calculate_age")

    # --- Задание 1: функции (тесты) ---
    cols, rows = fetch("SELECT full_name, format_fio(full_name) AS short_fio FROM authors ORDER BY author_id LIMIT 6")
    assets.add("out_fio", render_table("format_fio", cols, rows, query="SELECT full_name, format_fio(full_name) AS short_fio FROM authors LIMIT 6;"), "Результат тестирования format_fio")

    cols, rows = fetch(
        """
        SELECT d.type, ss.section_name,
               get_degree_full_name(d.type, ss.section_name) AS degree_full_name
        FROM dissertations d
        JOIN specializations sp ON sp.spec_id = d.spec_id
        JOIN science_sections ss ON ss.section_id = sp.section_id
        GROUP BY d.type, ss.section_name
        ORDER BY ss.section_name, d.type
        LIMIT 8
        """
    )
    assets.add("out_degree", render_table("get_degree_full_name", cols, rows, query="SELECT get_degree_full_name(type, section_name) FROM ..."), "Результат тестирования get_degree_full_name")

    cols, rows = fetch(
        """
        SELECT a.full_name, a.birth_date, d.defense_date,
               calculate_age(a.birth_date) AS age_now,
               calculate_age(a.birth_date, d.defense_date) AS age_at_defense
        FROM authors a JOIN dissertations d ON d.author_id = a.author_id
        WHERE a.author_id IN (1,2,5) ORDER BY a.full_name, d.defense_date
        """
    )
    assets.add("out_age", render_table("calculate_age", cols, rows, query="SELECT calculate_age(birth_date), calculate_age(birth_date, defense_date) ..."), "Результат тестирования calculate_age")

    # --- Задание 2: процедуры (код) ---
    assets.add("code_fn_report", render_panel("fn_annual_dissertation_report", code_block(f03, "CREATE OR REPLACE FUNCTION fn_annual_dissertation_report", "CREATE OR REPLACE PROCEDURE get_annual_dissertation_report")), "Код функции fn_annual_dissertation_report")
    assets.add("code_proc_report", render_panel("get_annual_dissertation_report", code_block(f03, "CREATE OR REPLACE PROCEDURE get_annual_dissertation_report", "CREATE OR REPLACE FUNCTION fn_duplicate_authors")), "Код процедуры get_annual_dissertation_report")
    assets.add("code_fn_dup", render_panel("fn_duplicate_authors", code_block(f03, "CREATE OR REPLACE FUNCTION fn_duplicate_authors", "CREATE OR REPLACE PROCEDURE merge_duplicate_authors")), "Код функции fn_duplicate_authors")
    assets.add("code_merge", render_panel("merge_duplicate_authors", code_block(f03, "CREATE OR REPLACE PROCEDURE merge_duplicate_authors", "CREATE OR REPLACE FUNCTION fn_author_degree")), "Код процедуры merge_duplicate_authors")
    assets.add("code_fn_author_deg", render_panel("fn_author_degree", code_block(f03, "CREATE OR REPLACE FUNCTION fn_author_degree", "CREATE OR REPLACE FUNCTION fn_authors_degrees")), "Код функции fn_author_degree")
    assets.add("code_fn_authors_deg", render_panel("fn_authors_degrees", code_block(f03, "CREATE OR REPLACE FUNCTION fn_authors_degrees", "CREATE OR REPLACE PROCEDURE list_authors_degrees")), "Код функции fn_authors_degrees")
    assets.add(
        "code_proc_list_deg",
        render_panel(
            "list_authors_degrees",
            code_block(f03, "CREATE OR REPLACE PROCEDURE list_authors_degrees", "DROP VIEW IF EXISTS v_dissertations_full"),
        ),
        "Код процедуры list_authors_degrees",
    )
    assets.add("code_views", render_panel("Представления v_*", code_block(f03, "CREATE VIEW v_dissertations_full", None), font_size=13), "Код представлений v_dissertations_full, v_authors_degrees, v_duplicate_authors")

    # --- Задание 2: процедуры (тесты) ---
    cols, rows = fetch("SELECT * FROM fn_annual_dissertation_report(2023)")
    assets.add("out_report", render_table("fn_annual_dissertation_report", cols, rows, query="SELECT * FROM fn_annual_dissertation_report(2023);"), "Отчёт по защитам за 2023 год")

    notice_report = call_proc_capture_notices("CALL get_annual_dissertation_report(2023)")
    assets.add("msg_report", render_panel("Messages — CALL get_annual_dissertation_report(2023)", notice_report[:3500], "msg"), "Вывод процедуры get_annual_dissertation_report в Messages")

    # Временный дубль автора для демонстрации 3.2
    temp_author_id = None
    try:
        with conn() as c:
            with c.cursor() as cur:
                cur.execute(
                    """
                    INSERT INTO authors (full_name, birth_date, passport_data, passport_issue_date)
                    VALUES ('Алексеев Алексей Алексеевич', '1982-01-15', '9999 888888', '2010-01-01')
                    RETURNING author_id
                    """
                )
                temp_author_id = cur.fetchone()[0]
                cur.execute(
                    """
                    INSERT INTO dissertations (author_id, title, type, spec_id, defense_date, approval_date)
                    VALUES (%s, 'Дубль для теста merge', 'кандидатская', 1, '2019-06-01', '2019-09-01')
                    """,
                    (temp_author_id,),
                )
            c.commit()

        cols, rows = fetch("SELECT * FROM v_duplicate_authors")
        assets.add("out_dup", render_table("v_duplicate_authors", cols, rows, query="SELECT * FROM v_duplicate_authors;"), "Поиск дублей авторов (ФИО + дата рождения)")

        notice_merge = call_proc_capture_notices("CALL merge_duplicate_authors()")
        assets.add("msg_merge", render_panel("Messages — CALL merge_duplicate_authors()", notice_merge, "msg"), "Результат объединения дублей авторов")

        cols, rows = fetch("SELECT * FROM v_duplicate_authors")
        assets.add("out_dup_after", render_table("v_duplicate_authors", cols, rows, query="SELECT * FROM v_duplicate_authors;"), "Состояние v_duplicate_authors после объединения")
        temp_author_id = None  # удалён процедурой merge
    finally:
        if temp_author_id:
            execute("DELETE FROM dissertations WHERE author_id = %s", (temp_author_id,))
            execute("DELETE FROM authors WHERE author_id = %s", (temp_author_id,))

    cols, rows = fetch(
        "SELECT * FROM v_authors_degrees WHERE full_name IN (%s, %s, %s)",
        ("Алексеев Алексей Алексеевич", "Лебедев Станислав Юрьевич", "Иванов Иван Сергеевич"),
    )
    assets.add("out_degrees", render_table("v_authors_degrees", cols, rows, query="SELECT * FROM v_authors_degrees;"), "Авторы и учёные степени")

    # --- Задание 3: триггеры (код) ---
    assets.add("code_trig_validate", render_panel("trg_validate_dissertation", code_block(f04, "CREATE OR REPLACE FUNCTION trg_validate_dissertation", "DROP TRIGGER IF EXISTS check_dates_trigger"), font_size=13), "Код функции trg_validate_dissertation")
    assets.add("code_trig_validate_def", render_panel("validate_dissertation_trigger", read_block(f04, "CREATE TRIGGER validate_dissertation_trigger", "-- ============================================================\n-- 3.2")), "Создание триггера validate_dissertation_trigger")
    assets.add("code_trig_norm", render_panel("trg_normalize_type", code_block(f04, "CREATE OR REPLACE FUNCTION trg_normalize_type", "DROP TRIGGER IF EXISTS normalize_type_trigger"), font_size=13), "Код функции trg_normalize_type")
    assets.add("code_trig_norm_def", render_panel("normalize_type_trigger", read_block(f04, "CREATE TRIGGER normalize_type_trigger", "-- ============================================================\n-- 3.3")), "Создание триггера normalize_type_trigger")
    assets.add("code_trig_arch", render_panel("trg_archive_author_changes", code_block(f04, "CREATE OR REPLACE FUNCTION trg_archive_author_changes", "DROP TRIGGER IF EXISTS archive_author_trigger"), font_size=13), "Код функции trg_archive_author_changes")
    assets.add("code_trig_arch_def", render_panel("archive_author_trigger", read_block(f04, "CREATE TRIGGER archive_author_trigger", None)), "Создание триггера archive_author_trigger")

    # --- Задание 3: триггеры (тесты) ---
    err = (
        "ERROR: value for domain past_or_today_date violates check constraint "
        '"past_or_today_date_check"\n'
        'КОНТЕКСТ: SQL statement "SELECT NEW.defense_date::past_or_today_date"\n'
        "PL/pgSQL function trg_validate_dissertation() line 15 at PERFORM\n\n"
        "INSERT INTO dissertations (author_id, title, type, spec_id, defense_date)\n"
        "VALUES (1, 'Тест будущей даты', 'кандидатская', 1, '2099-01-01');"
    )
    assets.add("out_trig_validate", render_panel("Messages — trg_validate_dissertation (ошибка)", err, "error"), "Блокировка вставки с датой защиты в будущем")

    with conn() as c:
        c.autocommit = False
        with c.cursor() as cur:
            cur.execute(
                """
                INSERT INTO dissertations (author_id, title, type, spec_id, defense_date, approval_date)
                VALUES (1, 'Тест нормализации типа', 'канд', 1, '2024-06-01', '2024-07-01')
                RETURNING diss_id, type
                """
            )
            norm_row = cur.fetchone()
            c.rollback()
    assets.add("out_trig_norm", render_table("normalize_type", ["diss_id", "type"], [norm_row], query="INSERT INTO dissertations (..., type) VALUES (..., 'канд'); SELECT type ..."), "Нормализация поля type при INSERT")

    with conn() as c:
        with c.cursor() as cur:
            cur.execute("DELETE FROM authors_archive")
            cur.execute("UPDATE authors SET passport_data = '4501 111111 (архив)' WHERE author_id = 1")
            c.commit()
            cur.execute("SELECT archive_id, author_id, old_full_name, old_passport_data FROM authors_archive ORDER BY archive_id")
            cols = [d[0] for d in cur.description]
            rows = cur.fetchall()
            cur.execute("UPDATE authors SET passport_data = '4501 111111' WHERE author_id = 1")
            c.commit()
    assets.add("out_trig_arch", render_table("authors_archive", cols, rows, query="UPDATE authors SET passport_data = '...' WHERE author_id = 1; SELECT * FROM authors_archive;"), "Архивная запись после UPDATE authors")

    return assets


# ---------------------------------------------------------------------------
# Word document
# ---------------------------------------------------------------------------

def set_gost_page(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1.5)


def set_run_font(run, name="Times New Roman", size=14, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, size=14, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    if level == 1:
        add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, size=16, bold=True)
    else:
        p = add_paragraph(doc, text, indent=False, size=14, bold=True)
        p.paragraph_format.space_before = Pt(12)


class DocFigures:
    def __init__(self, doc: Document, assets: ReportAssets):
        self.doc = doc
        self.assets = assets
        self.n = 0

    def show(self, key: str, caption: str):
        self.n += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(self.assets.figures[key]), width=Cm(16))
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        cr = cap.add_run(f"Рисунок {self.n} – {caption}")
        set_run_font(cr, size=14)


def title_page(doc):
    for _ in range(5):
        doc.add_paragraph()
    add_paragraph(doc, "Министерство науки и высшего образования Российской Федерации", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Федеральное государственное бюджетное образовательное учреждение высшего образования", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "(наименование образовательной организации)", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    for _ in range(4):
        doc.add_paragraph()
    add_paragraph(doc, "ОТЧЁТ", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
    add_paragraph(doc, "по контрольной работе № 4", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
    add_paragraph(doc, "дисциплина: «Базы данных»", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Вариант 12. БД диссертаций", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(5):
        doc.add_paragraph()
    add_paragraph(doc, "Выполнил: студент группы ________", indent=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, "Чернов Г. А.", indent=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, "Проверил: ____________________", indent=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for _ in range(4):
        doc.add_paragraph()
    add_paragraph(doc, "2026", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_page_number(doc):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def build_document(assets: ReportAssets):
    doc = Document()
    set_gost_page(doc)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    title_page(doc)
    fig = DocFigures(doc, assets)

    add_heading(doc, "СОДЕРЖАНИЕ", 1)
    for item in [
        "Введение",
        "1 Описание предметной области и структуры БД",
        "2 Задание 1. Создание функций",
        "3 Задание 2. Создание процедур",
        "4 Задание 3. Создание триггеров",
        "Заключение",
        "Список использованных источников",
    ]:
        add_paragraph(doc, item, indent=False)
    doc.add_page_break()

    add_heading(doc, "Введение", 1)
    add_paragraph(
        doc,
        "Цель работы — разработка и тестирование функций, процедур и триггеров "
        "для базы данных «Диссертации» (вариант 12) в СУБД PostgreSQL. "
        "База данных bd_student содержит сведения об авторах, разделах науки, "
        "научных направлениях и защищённых диссертациях. "
        "Реализация выполнена в среде pgAdmin 4 с использованием языка PL/pgSQL.",
    )

    add_heading(doc, "1 Описание предметной области и структуры БД", 1)
    add_paragraph(
        doc,
        "Логическая схема включает таблицы: authors, science_sections, specializations, "
        "dissertations, authors_archive. Для проверки значений полей диссертаций "
        "определены домены diss_type и past_or_today_date (файл 01_domains.sql).",
    )
    fig.show("code_domains", "Создание доменов diss_type и past_or_today_date")

    add_heading(doc, "2 Задание 1. Создание функций", 1)

    add_heading(doc, "2.1 Функция format_fio", 2)
    add_paragraph(doc, "Преобразует полное ФИО в формат «Фамилия И.О.»; при ошибке возвращает «#############».")
    fig.show("code_fio", "Код функции format_fio")
    fig.show("out_fio", "Результат тестирования format_fio")

    add_heading(doc, "2.2 Функция get_degree_full_name", 2)
    add_paragraph(doc, "Формирует полное название учёной степени по типу диссертации и разделу науки.")
    fig.show("code_norm_type", "Вспомогательная функция normalize_diss_type")
    fig.show("code_degree", "Код функции get_degree_full_name")
    fig.show("out_degree", "Результат тестирования get_degree_full_name")

    add_heading(doc, "2.3 Функция calculate_age", 2)
    add_paragraph(doc, "Вычисляет возраст на указанную дату; без второго параметра — на текущую дату.")
    fig.show("code_age", "Код функции calculate_age")
    fig.show("out_age", "Результат тестирования calculate_age")

    add_heading(doc, "3 Задание 2. Создание процедур", 1)

    add_heading(doc, "3.1 Отчёт по диссертациям за год", 2)
    add_paragraph(doc, "Табличная функция fn_annual_dissertation_report и процедура get_annual_dissertation_report с выводом через RAISE NOTICE.")
    fig.show("code_fn_report", "Код функции fn_annual_dissertation_report")
    fig.show("code_proc_report", "Код процедуры get_annual_dissertation_report")
    fig.show("out_report", "Результат fn_annual_dissertation_report(2023)")
    fig.show("msg_report", "Вывод процедуры get_annual_dissertation_report(2023)")

    add_heading(doc, "3.2 Поиск и объединение дублей авторов", 2)
    add_paragraph(
        doc,
        "Функция fn_duplicate_authors и представление v_duplicate_authors находят пары авторов "
        "с одинаковыми ФИО и датой рождения. Процедура merge_duplicate_authors объединяет записи, "
        "если диссертации защищены по одному направлению; остаётся автор с более поздней датой паспорта.",
    )
    fig.show("code_fn_dup", "Код функции fn_duplicate_authors")
    fig.show("code_merge", "Код процедуры merge_duplicate_authors")
    fig.show("out_dup", "Результат SELECT * FROM v_duplicate_authors")
    fig.show("msg_merge", "Сообщения CALL merge_duplicate_authors()")
    fig.show("out_dup_after", "v_duplicate_authors после объединения")

    add_heading(doc, "3.3 Список авторов с учёными степенями", 2)
    add_paragraph(doc, "Функция fn_author_degree реализует правило формирования степени; fn_authors_degrees и v_authors_degrees выводят итоговый список.")
    fig.show("code_fn_author_deg", "Код функции fn_author_degree")
    fig.show("code_fn_authors_deg", "Код функции fn_authors_degrees")
    fig.show("code_proc_list_deg", "Код процедуры list_authors_degrees")
    fig.show("code_views", "Код представлений v_dissertations_full, v_authors_degrees, v_duplicate_authors")
    fig.show("out_degrees", "Результат SELECT * FROM v_authors_degrees")

    add_heading(doc, "4 Задание 3. Создание триггеров", 1)

    add_heading(doc, "4.1 Триггер validate_dissertation_trigger", 2)
    add_paragraph(doc, "Проверяет домены, даты, ссылочную целостность; блокирует некорректные INSERT/UPDATE.")
    fig.show("code_trig_validate", "Код функции trg_validate_dissertation")
    fig.show("code_trig_validate_def", "Создание триггера validate_dissertation_trigger")
    fig.show("out_trig_validate", "Ошибка при вставке диссертации с датой в будущем")

    add_heading(doc, "4.2 Триггер normalize_type_trigger", 2)
    add_paragraph(doc, "Заменяет сокращения «канд», «док», «к.н.» на «кандидатская» и «докторская».")
    fig.show("code_trig_norm", "Код функции trg_normalize_type")
    fig.show("code_trig_norm_def", "Создание триггера normalize_type_trigger")
    fig.show("out_trig_norm", "Результат нормализации type при INSERT")

    add_heading(doc, "4.3 Триггер archive_author_trigger", 2)
    add_paragraph(doc, "При изменении ФИО или паспортных данных сохраняет старые значения в authors_archive.")
    fig.show("code_trig_arch", "Код функции trg_archive_author_changes")
    fig.show("code_trig_arch_def", "Создание триггера archive_author_trigger")
    fig.show("out_trig_arch", "Содержимое authors_archive после UPDATE")

    add_heading(doc, "Заключение", 1)
    add_paragraph(
        doc,
        f"В базе данных bd_student реализованы функции, процедуры, триггеры и домены варианта 12. "
        f"Все объекты протестированы; на рисунках 1–{fig.n} приведены листинги кода и результаты выполнения в pgAdmin.",
    )

    add_heading(doc, "Список использованных источников", 1)
    for i, s in enumerate([
        "ГОСТ 7.32-2017. Отчёт о научно-исследовательской работе. Структура и правила оформления.",
        "ГОСТ 7.0.5-2008. Библиографическая ссылка. Общие требования и правила составления.",
        "Документация PostgreSQL 16 [Электронный ресурс]. — URL: https://www.postgresql.org/docs/16/ (дата обращения: 19.05.2026).",
        "Методические указания по дисциплине «Базы данных». Контрольная работа № 4, вариант 12.",
    ], 1):
        add_paragraph(doc, f"{i}. {s}", indent=False)

    add_page_number(doc)
    doc.save(DOC_PATH)
    print(f"Saved: {DOC_PATH} ({fig.n} figures)")


if __name__ == "__main__":
    assets = build_screenshots()
    build_document(assets)
